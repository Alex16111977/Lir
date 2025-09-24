"""
Генератор промптів для Grok-4 v8.4
Оптимізована версія для обробки всіх JSON файлів з папки data
Створює DOCX файли в структурі Ani/grok/{a2,b1,thematic}
ИСПРАВЛЕНО: теперь используются ВСЕ 12 слов из JSON
"""

import json
import sys
import subprocess
import importlib.util
from pathlib import Path
from datetime import datetime

# Перевірка та встановлення python-docx
def ensure_docx_installed():
    """Перевірка та встановлення python-docx"""
    if importlib.util.find_spec("docx") is None:
        print("[!] Встановлюю python-docx...")
        venv_pip = r"F:\AiKlientBank\Lir\.venv\Scripts\pip.exe"
        result = subprocess.run(
            [venv_pip, "install", "python-docx"],
            capture_output=True,
            text=True
        )
        if result.returncode == 0:
            print("[OK] python-docx встановлено")
        else:
            print(f"[ERROR] Не вдалося встановити: {result.stderr}")
            sys.exit(1)

ensure_docx_installed()

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Константи
DATA_DIR = Path(r"F:\AiKlientBank\Lir\data")
OUTPUT_DIR = Path(r"F:\AiKlientBank\Lir\Ani\grok")

def load_json_data(json_path):
    """Завантаження JSON даних з обробкою помилок"""
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        print(f"[OK] JSON завантажено: {json_path.name}")
        print(f"[+] Заголовок: {data.get('title', '')}")
        print(f"[+] Кількість слів: {len(data.get('vocabulary', []))}")
        return data
    except json.JSONDecodeError as e:
        print(f"[ERROR] Помилка JSON: {e} в {json_path}")
        return None
    except Exception as e:
        print(f"[ERROR] Помилка завантаження: {e} в {json_path}")
        return None

class GrokDocxCreator:
    """Генератор промпта з готовою історією"""
    
    def __init__(self, json_data, json_filename="урок"):
        self.data = json_data
        self.filename = json_filename.replace('.json', '')
        self.doc = Document()
        self.vocabulary = self.data.get('vocabulary', [])[:12]  # Все 12 слів з JSON
        self.setup_document()
    
    def setup_document(self):
        """Налаштування стилів документа"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        style = self.doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        print("[OK] Документ створено")
    
    def format_transcription(self, transcription):
        """Форматування транскрипції"""
        if not transcription:
            return ""
        formatted = transcription.replace('[', '').replace(']', '')
        formatted = formatted.replace(' ', '-')
        return formatted
    
    def add_header(self):
        """Додавання заголовка - ВІДКЛЮЧЕНО"""
        # Заголовок видалено за запитом користувача
        pass
    
    def create_theatrical_story(self):
        """Створення унікального театрального тексту з ВСІМА 12 німецькими словами"""
        
        # Отримуємо дані з JSON
        title = self.data.get('title', '').replace('🎭', '').replace('🏰', '').replace('👑', '').strip()
        emotions = self.data.get('emotions', [])
        dialogues = self.data.get('dialogues', [])
        
        # Створюємо унікальний текст для кожного файлу
        story_text = f"### {title.upper()}\n\n"
        
        # Вступний параграф з емоціями
        if emotions:
            clean_emotions = [e.replace('⚡', '').replace('👑', '').replace('🎭', '').strip() 
                            for e in emotions]
            story_text += f"Эмоции сцены: {', '.join(clean_emotions)}\n\n"
        
        # АКТ 1: ВСТУП (слова 1-4)
        story_text += "**АКТ 1: ВСТУПЛЕНИЕ**\n\n"
        story_text += "Тронный зал. Король Лир входит.\n\n"
        
        for i, word in enumerate(self.vocabulary[:4]):
            german = word.get('german', '')
            transcription = self.format_transcription(word.get('transcription', ''))
            translation = word.get('translation', '')
            character_voice = word.get('character_voice', {})
            character = character_voice.get('character', 'Рассказчик')
            voice_german = character_voice.get('german', '')
            voice_russian = character_voice.get('russian', '')
            
            if german and transcription and translation:
                story_text += f"{character}: "
                if voice_german:
                    story_text += f'"{voice_german}"\n'
                    story_text += f'(**{german.upper()}** [{transcription}] ... {translation})\n'
                    if voice_russian:
                        story_text += f'({voice_russian})\n'
                else:
                    story_text += f"В этой сцене важное слово **{german.upper()}** [{transcription}] ... "
                    story_text += f"({translation})\n"
                
                story_text += "\n"
        
        story_text += "*** [драматическая пауза] ***\n\n"
        
        # АКТ 2: РАЗВИТИЕ (слова 5-8)
        story_text += "**АКТ 2: РАЗВИТИЕ ДЕЙСТВИЯ**\n\n"
        
        for i, word in enumerate(self.vocabulary[4:8]):
            german = word.get('german', '')
            transcription = self.format_transcription(word.get('transcription', ''))
            translation = word.get('translation', '')
            character_voice = word.get('character_voice', {})
            character = character_voice.get('character', 'Рассказчик')
            voice_german = character_voice.get('german', '')
            voice_russian = character_voice.get('russian', '')
            
            if german and transcription and translation:
                story_text += f"{character}: "
                if voice_german:
                    story_text += f'"{voice_german}"\n'
                    story_text += f'(**{german.upper()}** [{transcription}] ... {translation})\n'
                    if voice_russian:
                        story_text += f'({voice_russian})\n'
                else:
                    story_text += f"Ключевое слово момента **{german.upper()}** [{transcription}] ... "
                    story_text += f"({translation})\n"
                
                story_text += "\n"
                
                # Пауза после каждых 2 слов
                if (i + 1) % 2 == 0:
                    story_text += "*** [пауза для осмысления] ***\n\n"
        
        # АКТ 3: КУЛЬМИНАЦИЯ (слова 9-10)
        story_text += "**АКТ 3: ЭМОЦИОНАЛЬНЫЙ ПИК**\n\n"
        
        # Используем диалоги, если есть
        if dialogues:
            for d in dialogues[:2]:
                character = d.get('character', '').replace('🎭', '').replace('👑', '').strip()
                emotion = d.get('emotion', '').replace('🎭', '').replace('👑', '').strip()
                dialog_german = d.get('german', '')
                dialog_russian = d.get('russian', '')
                
                if character and dialog_german:
                    story_text += f"{character}: [{emotion}]\n"
                    story_text += f'"{dialog_german}"\n'
                    story_text += f'({dialog_russian})\n\n'
        
        # Слова 9-10
        for word in self.vocabulary[8:10]:
            german = word.get('german', '')
            transcription = self.format_transcription(word.get('transcription', ''))
            translation = word.get('translation', '')
            
            if german:
                story_text += f"Кульминационное слово: **{german.upper()}** [{transcription}] ... ({translation})\n"
        
        story_text += "\n*** [напряжённая пауза] ***\n\n"
        
        # АКТ 4: ФИНАЛ (слова 11-12)
        story_text += "**АКТ 4: ФИНАЛ**\n\n"
        
        for word in self.vocabulary[10:12]:
            german = word.get('german', '')
            transcription = self.format_transcription(word.get('transcription', ''))
            translation = word.get('translation', '')
            character_voice = word.get('character_voice', {})
            character = character_voice.get('character', 'Лир')
            
            if german:
                story_text += f"{character}: "
                story_text += f"В финале звучит **{german.upper()}** [{transcription}] ... ({translation})\n"
        
        story_text += "\n*** Занавес падает ***\n"
        
        # Добавляем сводку всех слов в конце для проверки
        story_text += "\n**СЛОВАРЬ УРОКА (все 12 слов):**\n"
        for i, word in enumerate(self.vocabulary[:12]):
            german = word.get('german', '')
            if german:
                story_text += f"{i+1}. {german.upper()}\n"
        
        return story_text
    
    def add_theatrical_prompt(self):
        """Створення промпту з чітким розділенням інструкцій і тексту - РУССКИЙ ЯЗЫК"""
        
        # Дані з JSON
        title = self.data.get('title', '').replace('🎭', '').replace('🏰', '').strip()
        dialogues = self.data.get('dialogues', [])
        memory = self.data.get('memory_trick', {})
        if isinstance(memory, str):
            memory = {'master_trick': memory}
        
        # 1. ЧІТКА ІНСТРУКЦІЯ НА ПОЧАТКУ - НА РУССКОМ
        self.doc.add_heading('# ИНСТРУКЦИИ ДЛЯ GROK (НЕ ЧИТАТЬ ВСЛУХ)', level=1)
        critical = self.doc.add_paragraph()
        critical.add_run('ВАЖНО: Читай вслух ТОЛЬКО содержимое раздела "ТЕКСТ ДЛЯ ОЗВУЧИВАНИЯ". Всё остальное - это инструкции для тебя, НЕ читай их вслух.').bold = True
        critical.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        
        # 2. ЛЕГЕНДА СИМВОЛІВ З ВИРАЗНИМИ ІНСТРУКЦІЯМИ - НА РУССКОМ
        self.doc.add_heading('# ЛЕГЕНДА СИМВОЛОВ (НЕ ЧИТАЙ ВСЛУХ)', level=2)
        legend = self.doc.add_paragraph()
        legend.add_run('... = МОЛЧАНИЕ 2 СЕКУНДЫ (НЕ ПРОИЗНОСИ СЛОВО "ПАУЗА" ИЛИ "ТОЧКИ")\n')
        legend.add_run('*** = МОЛЧАНИЕ 3 СЕКУНДЫ (НЕ ПРОИЗНОСИ СИМВОЛЫ)\n')
        legend.add_run('[текст] = ИНСТРУКЦИЯ ДЛЯ ИНТОНАЦИИ (НЕ ЧИТАЙ ВСЛУХ)\n')
        legend.add_run('Имя: = СМЕНИ ГОЛОС НА ПЕРСОНАЖА (НЕ ПРОИЗНОСИ ИМЯ С ДВОЕТОЧИЕМ)\n')
        legend.runs[0].font.color.rgb = RGBColor(128, 0, 0)
        
        # 3. ГОЛОВНА ІНСТРУКЦІЯ З ЯСНИМИ ВКАЗІВКАМИ - НА РУССКОМ
        self.doc.add_heading('# ГЛАВНАЯ ИНСТРУКЦИЯ (НЕ ЧИТАЙ ВСЛУХ)', level=2)
        instruction = self.doc.add_paragraph()
        instruction.add_run('Ты — актёр-рассказчик. Озвучь ТОЛЬКО текст в разделе "ТЕКСТ ДЛЯ ОЗВУЧИВАНИЯ", меняя голоса для персонажей. Когда видишь символы ... или ***, сделай паузу, но НЕ ПРОИЗНОСИ "пауза" или "точки". Транскрипцию в скобках используй для правильного произношения немецких слов.')
        
        # 4. ПРАВИЛА ОЗВУЧУВАННЯ - НА РУССКОМ
        self.doc.add_heading('## ПРАВИЛА ОЗВУЧИВАНИЯ (НЕ ЧИТАЙ ВСЛУХ)', level=3)
        rules = [
            '• Немецкие слова (ЗАГЛАВНЫМИ) - произноси МЕДЛЕННО по транскрипции',
            '• После немецкого слова - МОЛЧИ 2 секунды (НЕ ГОВОРИ "пауза")',
            '• Перевод в круглых скобках - произноси после молчания',
            '• При указании персонажа - смени голос, НЕ читай имя',
            '• В эмоциональном пике - максимальная драматичность',
            '• ВСЕГО В УРОКЕ 12 СЛОВ - озвучь ВСЕ'
        ]
        for rule in rules:
            self.doc.add_paragraph(rule)
        
        # 5. ПЕРСОНАЖІ З ОПИСОМ ГОЛОСІВ - НА РУССКОМ
        self.doc.add_heading('## ГОЛОСА ПЕРСОНАЖЕЙ (НЕ ЧИТАЙ ВСЛУХ)', level=3)
        characters = [
            'Король Лир: глубокий, усталый голос старого правителя',
            'Корделия: нежный, искренний голос младшей дочери',
            'Гонерилья: холодный, расчётливый голос старшей дочери',
            'Регана: льстивый, но жестокий голос средней дочери',
            'Глостер: официальный, церемониальный голос старого графа',
            'Эдмунд: амбициозный, дерзкий голос молодого интригана',
            'Кент: верный, решительный голос советника',
            'Рассказчик: нейтральный, драматичный голос'
        ]
        char_para = self.doc.add_paragraph()
        for char in characters:
            char_para.add_run(f"{char}\n")
        
        # 6. ЧІТКО ВІДОКРЕМЛЕНИЙ РОЗДІЛ З ТЕКСТОМ ДЛЯ ОЗВУЧУВАННЯ - НА РУССКОМ
        self.doc.add_page_break()
        self.doc.add_heading('# ТЕКСТ ДЛЯ ОЗВУЧИВАНИЯ (ЧИТАЙ ТОЛЬКО ЭТОТ РАЗДЕЛ)', level=1)
        
        # Додаємо театральний текст
        story_text = self.create_theatrical_story()
        for line in story_text.split('\n'):
            if line.strip():
                if line.startswith('###'):
                    h = self.doc.add_heading(line.replace('###', '').strip(), level=3)
                    h.runs[0].font.color.rgb = RGBColor(139, 0, 0)
                elif '**' in line:
                    # Обробка жирного тексту
                    p = self.doc.add_paragraph()
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 1:  # Жирний текст
                            run = p.add_run(part)
                            run.bold = True
                            if part.isupper() or '[' in part:  # Німецькі слова
                                run.font.color.rgb = RGBColor(0, 0, 200)
                        else:
                            p.add_run(part)
                else:
                    self.doc.add_paragraph(line)
        
        # 7. ІНТЕРАКТИВНІ МОМЕНТИ - НА РУССКОМ (со всеми 12 словами)
        self.doc.add_heading('## ИНТЕРАКТИВНЫЕ МОМЕНТЫ (ВСТАВЛЯЙ В ИСТОРИЮ)', level=3)
        interactive_para = self.doc.add_paragraph()
        interactive_para.add_run('НЕ ЧИТАЙ ЭТОТ ЗАГОЛОВОК ВСЛУХ!\n').bold = True
        
        # Используем больше слов для интерактива
        interactive_words = [0, 3, 6, 9]  # Слова 1, 4, 7, 10
        for idx in interactive_words:
            if idx < len(self.vocabulary):
                german = self.vocabulary[idx].get('german', '')
                if german:
                    if idx == 0:
                        interactive_para.add_run(f'• После слова {german}: "Повтори за мной: {german}! ... Покажи жест рукой!"\n')
                    elif idx == 3:
                        interactive_para.add_run(f'• После слова {german}: "Какая ассоциация у тебя с этим словом? ... Интересно!"\n')
                    elif idx == 6:
                        interactive_para.add_run(f'• После слова {german}: "Попробуй составить предложение с этим словом! ..."\n')
                    elif idx == 9:
                        interactive_para.add_run(f'• После слова {german}: "Это слово из кульминации! Покажи эмоцию!"\n')
        
        interactive_para.add_run('• В конце истории: "Какие слова ты запомнил? Давай повторим все 12 слов вместе!"\n')
        
        # 8. МНЕМОНІЧНИЙ ПРИЙОМ - НА РУССКОМ
        if memory.get('master_trick'):
            self.doc.add_heading('## МНЕМОНИЧЕСКИЙ ПРИЁМ (НЕ ЧИТАЙ ВСЛУХ)', level=3)
            mem_para = self.doc.add_paragraph()
            mem_para.add_run(f'{memory["master_trick"]}')
        
        # 9. ПОЧАТОК І КІНЕЦЬ УРОКУ - НА РУССКОМ
        self.doc.add_heading('## НАЧАЛО И КОНЕЦ УРОКА', level=3)
        start_end = self.doc.add_paragraph()
        start_end.add_run('НЕ ЧИТАЙ ЭТОТ ЗАГОЛОВОК ВСЛУХ!\n').bold = True
        start_end.add_run('\nНачни так:\n').bold = True
        start_end.add_run('"Добро пожаловать в театр Короля Лира! Сегодня ты выучишь 12 немецких слов через драматическую историю в 4 актах. Слушай, повторяй, показывай жесты! Готов? ... Занавес открывается..."\n\n')
        
        start_end.add_run('Закончи так:\n').bold = True
        start_end.add_run('"Браво! Ты прошёл через все 4 акта и выучил 12 новых немецких слов! Какие из них тебе больше всего запомнились? ... Отлично! До следующей встречи в театре Короля Лира!"')
        
        # 10. ФІНАЛЬНЕ НАГАДУВАННЯ - НА РУССКОМ
        self.doc.add_paragraph()
        reminder = self.doc.add_paragraph()
        run = reminder.add_run(
            'ПОМНИ: Символы ... и *** означают МОЛЧАНИЕ! ' +
            'НЕ ПРОИЗНОСИ ИХ или слова "пауза", "молчание" и т.д.! ' +
            'Читай ТОЛЬКО текст из раздела "ТЕКСТ ДЛЯ ОЗВУЧИВАНИЯ"! ' +
            'В уроке должно прозвучать ВСЕ 12 НЕМЕЦКИХ СЛОВ!'
        )
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 0, 0)
    
    def save(self, output_dir, category):
        """Збереження документа з обробкою помилок"""
        try:
            # Створення директорії категорії, якщо її ще немає
            category_dir = output_dir / category
            category_dir.mkdir(parents=True, exist_ok=True)
            
            # Очищення імені файлу від спеціальних символів
            clean_name = self.filename.replace(':', '').replace('/', '_').replace(' ', '_')
            output_file = category_dir / f"GROK_PROMPT_{clean_name}.docx"
            
            # Збереження файлу
            self.doc.save(output_file)
            print(f"[SUCCESS] Документ збережено: {output_file}")
            print(f"[INFO] Розмір: {output_file.stat().st_size / 1024:.2f} KB")
            return output_file
        except Exception as e:
            print(f"[ERROR] Помилка збереження: {e}")
            return None

def process_json_file(json_path, output_dir, category):
    """Обробка одного JSON файлу"""
    try:
        data = load_json_data(json_path)
        if not data:
            print(f"[SKIP] Пропускаю файл через помилку: {json_path.name}")
            return None
        
        creator = GrokDocxCreator(data, json_path.stem)
        # creator.add_header()  # Заголовок відключено
        creator.add_theatrical_prompt()
        
        return creator.save(output_dir, category)
    except Exception as e:
        print(f"[ERROR] Помилка обробки {json_path.name}: {e}")
        return None

def process_all_json_files():
    """Обробка всіх JSON файлів з папки data"""
    results = []
    
    # Перевірка категорій
    categories = [d for d in DATA_DIR.iterdir() if d.is_dir()]
    print(f"\n[INFO] Знайдено {len(categories)} категорій: {', '.join([c.name for c in categories])}")
    
    # Обробка всіх категорій
    for category_dir in categories:
        category = category_dir.name
        print(f"\n{'='*60}")
        print(f"[PROCESS] Обробка категорії: {category}")
        print('='*60)
        
        # Створення вихідної директорії для категорії
        (OUTPUT_DIR / category).mkdir(parents=True, exist_ok=True)
        
        # Обробка всіх JSON файлів у категорії
        json_files = list(category_dir.glob("*.json"))
        print(f"[INFO] Знайдено {len(json_files)} JSON файлів у категорії {category}")
        
        for json_file in json_files:
            print(f"\n{'-'*40}")
            print(f"[PROCESS] Обробка: {json_file.name}")
            print(f"{'-'*40}")
            
            result = process_json_file(json_file, OUTPUT_DIR, category)
            if result:
                results.append((category, json_file.name, result))
    
    return results

def main():
    """Головна функція"""
    print("="*60)
    print("Grok-4 Theatrical Prompt Generator v8.4 (RUSSIAN + ALL 12 WORDS)")
    print("Обробка всіх JSON файлів з папки data")
    print("ИСПРАВЛЕНО: используются ВСЕ 12 слов")
    print("="*60)
    
    # Переконатися, що вихідні директорії існують
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for category in ['a2', 'b1', 'thematic']:
        (OUTPUT_DIR / category).mkdir(parents=True, exist_ok=True)
    
    print(f"[CONFIG] Шлях до JSON файлів: {DATA_DIR}")
    print(f"[CONFIG] Шлях для збереження: {OUTPUT_DIR}")
    
    # Обробка всіх JSON файлів
    results = process_all_json_files()
    
    # Статистика
    print("\n" + "="*60)
    print(f"[COMPLETE] Створено {len(results)} документів")
    print(f"[STATS] Категорії: {len(set([r[0] for r in results]))}")
    for category in sorted(set([r[0] for r in results])):
        category_count = sum(1 for r in results if r[0] == category)
        print(f"  [+] {category}: {category_count} файлів")
    print("="*60)
    
    print("\n[INFO] Все 12 слов теперь используются в каждом уроке:")
    print("  - АКТ 1: слова 1-4")
    print("  - АКТ 2: слова 5-8")
    print("  - АКТ 3: слова 9-10")
    print("  - АКТ 4: слова 11-12")
    print("  - СЛОВАРЬ: все 12 слов в конце")
    
    print("\n[FINISH] Обробку завершено!")
    print("[SUCCESS] Всі промпти на РУССКОМ ЯЗЫКЕ з усіма 12 словами!")

if __name__ == "__main__":
    main()
